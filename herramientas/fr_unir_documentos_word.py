"""
Herramienta: FR_unir_documentos_word
Descripcion: Unifica todos los archivos Word (.doc y .docx) de un directorio en un único documento. Los títulos de cada archivo aparecen en MAYÚSCULAS, NEGRITA y SUBRAYADO, seguidos de su contenido en orden alfabético.
Generada automaticamente el 2026-02-21 18:41:09
Autor: Sistema de Generacion Automatica
"""

from mcp.types import Tool, TextContent

# Definicion de la herramienta
HERRAMIENTA = Tool(
    name="FR_unir_documentos_word",
    description="Unifica todos los archivos Word (.doc y .docx) de un directorio en un único documento. Los títulos de cada archivo aparecen en MAYÚSCULAS, NEGRITA y SUBRAYADO, seguidos de su contenido en orden alfabético.",
    inputSchema={
        "type": "object",
        "properties": {
                "ruta_carpeta": {
                        "type": "string",
                        "description": "Ruta completa del directorio que contiene los archivos Word a unir. Ej: C:\\Users\\usuario\\Documents\\carpeta"
                },
                "nombre_salida": {
                        "type": "string",
                        "description": "Nombre del archivo de salida (sin extensión). Por defecto: DOCUMENTO_UNIFICADO"
                }
        },
        "required": [
                "ruta_carpeta"
        ]
}
)

# Funcion de ejecucion
async def ejecutar(argumentos: dict) -> list[TextContent]:
    """Ejecuta la herramienta FR_unir_documentos_word"""
    
import os
import subprocess
import sys
from pathlib import Path

def unir_documentos_word(ruta_carpeta, nombre_salida="DOCUMENTO_UNIFICADO"):
    """
    Unifica archivos Word de un directorio en un documento único.
    """
    try:
        import win32com.client
    except:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "pywin32"])
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import win32com.client
    
    from docx import Document
    from docx.shared import Pt
    
    # Validar ruta
    if not os.path.isdir(ruta_carpeta):
        return {"error": f"La ruta no existe: {ruta_carpeta}"}
    
    # Obtener archivos ordenados
    archivos = sorted([f for f in os.listdir(ruta_carpeta) if f.endswith(('.doc', '.docx'))])
    
    if not archivos:
        return {"error": f"No se encontraron archivos Word en: {ruta_carpeta}"}
    
    # Iniciar Word
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    
    # Crear documento nuevo
    doc_unificado = Document()
    
    procesados = []
    errores = []
    
    for archivo in archivos:
        ruta_archivo = os.path.join(ruta_carpeta, archivo)
        
        try:
            # Abrir con Word COM
            doc_word = word.Documents.Open(os.path.abspath(ruta_archivo), ReadOnly=True)
            
            # Agregar título en mayúsculas, negrita y subrayado
            titulo = doc_unificado.add_paragraph()
            run = titulo.add_run(archivo.upper())
            run.bold = True
            run.underline = True
            run.font.size = Pt(12)
            
            # Extraer texto
            for parrafo in doc_word.Paragraphs:
                if parrafo.Range.Text.strip():
                    doc_unificado.add_paragraph(parrafo.Range.Text)
            
            # Extraer tablas
            for tabla_word in doc_word.Tables:
                rows = len(tabla_word.Rows)
                cols = len(tabla_word.Columns)
                tabla_nueva = doc_unificado.add_table(rows=rows, cols=cols)
                
                for i, fila in enumerate(tabla_word.Rows):
                    for j, celda in enumerate(fila.Cells):
                        tabla_nueva.rows[i].cells[j].text = celda.Range.Text
            
            # Separador
            doc_unificado.add_paragraph()
            
            doc_word.Close(False)
            procesados.append(archivo)
            
        except Exception as e:
            errores.append({"archivo": archivo, "error": str(e)})
    
    # Cerrar Word
    word.Quit()
    
    # Guardar documento
    ruta_salida = os.path.join(ruta_carpeta, f"{nombre_salida}.docx")
    doc_unificado.save(ruta_salida)
    
    return {
        "exito": True,
        "archivo_creado": ruta_salida,
        "archivos_procesados": len(procesados),
        "archivos": procesados,
        "errores": errores if errores else None,
        "mensaje": f"✅ Documento unificado creado: {nombre_salida}.docx"
    }
    
    return [TextContent(type="text", text=resultado)]
